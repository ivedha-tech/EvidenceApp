import asyncio
from playwright.async_api import async_playwright
from docx import Document
from datetime import datetime
import os
from PIL import Image, ImageDraw, ImageFont
import sys

async def add_timestamp_to_image(image_path):
    # Open the image
    img = Image.open(image_path)
    draw = ImageDraw.Draw(img)
    
    # Add timestamp
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    # You might need to adjust the font path based on your system
    try:
        font = ImageFont.truetype("arial.ttf", 24)
    except:
        font = ImageFont.load_default()
    
    # Add text to image
    draw.text((10, 10), timestamp, fill="red", font=font)
    img.save(image_path)

async def process_asn(asn, url):
    async with async_playwright() as p:
        browser = await p.chromium.launch()
        page = await browser.new_page()
        
        try:
            # Navigate to the URL
            await page.goto(url)
            
            # Create screenshots directory if it doesn't exist
            if not os.path.exists("screenshots"):
                os.makedirs("screenshots")
            
            # Take screenshot
            screenshot_path = f"screenshots/{asn}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png"
            await page.screenshot(path=screenshot_path)
            
            # Add timestamp to the screenshot
            await add_timestamp_to_image(screenshot_path)
            
            # Create Word document
            doc = Document()
            doc.add_heading(f'ASN: {asn}', 0)
            doc.add_paragraph(f'URL: {url}')
            doc.add_paragraph(f'Timestamp: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_picture(screenshot_path)
            
            # Save Word document
            doc.save(f"{asn}_report.docx")
            
        except Exception as e:
            print(f"Error processing ASN {asn}: {str(e)}")
        finally:
            await browser.close()

async def main():
    if len(sys.argv) != 3:
        print("Usage: python asn_screenshot.py <comma_separated_asns> <url>")
        sys.exit(1)
    
    asns = [asn.strip() for asn in sys.argv[1].split(',')]
    url = sys.argv[2]
    
    for asn in asns:
        print(f"Processing ASN: {asn}")
        await process_asn(asn, url)
        print(f"Completed processing ASN: {asn}")

if __name__ == "__main__":
    asyncio.run(main()) 